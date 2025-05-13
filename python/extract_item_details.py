import json
import os
import shutil
from docx import Document
import re

# --- Configuration ---
processed_data_json_path = "/home/ubuntu/site_bazar/processed_inventory_data.json"
zip_extracted_base_dir = "/home/ubuntu/inventory_check"
site_images_base_dir = "/home/ubuntu/site_bazar/assets/img"
output_full_data_path = "/home/ubuntu/site_bazar/full_item_data.json"
output_extraction_report_path = "/home/ubuntu/site_bazar/extraction_report.txt"

allowed_image_extensions = (".jpg", ".jpeg", ".png", ".webp", ".avif")

# --- Helper function to extract descriptions ---
def extract_descriptions_from_docx(docx_path):
    pt_desc_str = ""
    en_desc_str = ""
    paragraphs_text = []
    try:
        document = Document(docx_path)
        if not document.paragraphs:
            return "", "" # Empty document

        for para in document.paragraphs:
            paragraphs_text.append(para.text.strip())
        
        full_text = "\n".join(paragraphs_text)

        # Try to find explicit language markers
        pt_marker_regex = r"^(Portugu[êe]s|PT)[:\s-]*\n?"
        en_marker_regex = r"^(Ingl[êe]s|English|EN)[:\s-]*\n?"
        
        pt_sections = []
        en_sections = []
        current_lang = None # 'PT', 'EN', or None
        
        temp_buffer = []
        
        for para_text in paragraphs_text:
            is_pt_marker = re.match(pt_marker_regex, para_text, re.IGNORECASE)
            is_en_marker = re.match(en_marker_regex, para_text, re.IGNORECASE)

            if is_pt_marker:
                if current_lang == 'EN' and temp_buffer:
                    en_sections.append("\n".join(temp_buffer))
                elif current_lang == 'PT' and temp_buffer: # Consecutive PT markers, append previous buffer
                    pt_sections.append("\n".join(temp_buffer))
                temp_buffer = [re.sub(pt_marker_regex, "", para_text, flags=re.IGNORECASE).strip()] # Remove marker from text
                current_lang = 'PT'
                continue
            elif is_en_marker:
                if current_lang == 'PT' and temp_buffer:
                    pt_sections.append("\n".join(temp_buffer))
                elif current_lang == 'EN' and temp_buffer: # Consecutive EN markers
                    en_sections.append("\n".join(temp_buffer))
                temp_buffer = [re.sub(en_marker_regex, "", para_text, flags=re.IGNORECASE).strip()] # Remove marker
                current_lang = 'EN'
                continue
            
            if para_text: # Only add non-empty paragraphs to the current buffer
                temp_buffer.append(para_text)
        
        # Append the last buffer
        if current_lang == 'PT' and temp_buffer:
            pt_sections.append("\n".join(temp_buffer))
        elif current_lang == 'EN' and temp_buffer:
            en_sections.append("\n".join(temp_buffer))
        
        pt_desc_str = "\n".join(filter(None, pt_sections)).strip()
        en_desc_str = "\n".join(filter(None, en_sections)).strip()

        # Fallback if markers are not clear or only one language block was found this way
        if not pt_desc_str and not en_desc_str and paragraphs_text:
            # If no markers at all, assume all is PT
            pt_desc_str = "\n".join(filter(None, paragraphs_text)).strip()
        elif not en_desc_str and pt_desc_str: # If only PT found, use PT for EN as per todo.md
            en_desc_str = pt_desc_str 
        elif not pt_desc_str and en_desc_str: # If only EN found (less likely but possible)
             pt_desc_str = en_desc_str # Fallback PT to EN

        return pt_desc_str, en_desc_str

    except Exception as e:
        return f"Error reading DOCX: {str(e)}", f"Error reading DOCX: {str(e)}"

# --- Main Script ---
if not os.path.exists(processed_data_json_path):
    with open(output_extraction_report_path, 'w', encoding='utf-8') as r_file:
        r_file.write(f"Erro Crítico: Arquivo de dados processados não encontrado em {processed_data_json_path}\n")
    print(f"CRITICAL ERROR: Processed data JSON not found. Report: {output_extraction_report_path}")
    exit()

with open(processed_data_json_path, 'r', encoding='utf-8') as f:
    mapped_items = json.load(f)

full_items_data = []
extraction_issues = []

if not os.path.exists(site_images_base_dir):
    os.makedirs(site_images_base_dir)

for item in mapped_items:
    item_id = item.get("id")
    folder_name = item.get("folder_name")
    item_title_for_report = item.get('title_pt', item.get('csv_item_name_original', item_id)) 
    original_csv_row = item.get("original_csv_row", "N/A")

    if not folder_name:
        extraction_issues.append(f"Item (Row {original_csv_row}, ID {item_id}): Missing folder_name in processed data.")
        continue

    source_item_folder_path = os.path.join(zip_extracted_base_dir, folder_name)
    target_item_image_dir = os.path.join(site_images_base_dir, folder_name)

    item_data = item.copy()
    item_data["images"] = []
    item_data["description_pt"] = ""
    item_data["description_en"] = ""
    item_data["description_en_is_fallback"] = False # Initialize this flag

    # 1. Copy Images
    if not os.path.exists(source_item_folder_path):
        extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Pasta de origem não encontrada em {source_item_folder_path}.")
    else:
        if not os.path.exists(target_item_image_dir):
            try:
                os.makedirs(target_item_image_dir)
            except OSError as e:
                extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}): Erro ao criar diretório de imagem {target_item_image_dir}: {str(e)}.")
                continue 
        
        images_found_for_item = False
        for filename in os.listdir(source_item_folder_path):
            if filename.lower().endswith(allowed_image_extensions):
                source_image_path = os.path.join(source_item_folder_path, filename)
                target_image_path = os.path.join(target_item_image_dir, filename)
                try:
                    shutil.copy2(source_image_path, target_image_path)
                    item_data["images"].append(f"assets/img/{folder_name}/{filename}")
                    images_found_for_item = True
                except Exception as e:
                    extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}): Erro ao copiar imagem {filename}: {str(e)}.")
        if not images_found_for_item:
            extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Nenhuma imagem ({', '.join(allowed_image_extensions)}) encontrada na pasta de origem.")

    # 2. Extract Descriptions from DOCX
    docx_file_path = os.path.join(source_item_folder_path, "description.docx")
    if not os.path.exists(docx_file_path):
        extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Arquivo description.docx não encontrado.")
        item_data["description_pt"] = item.get("description_pt_raw_csv", "Descrição em Português não disponível.")
        item_data["description_en"] = item_data["description_pt"] # Fallback EN to PT
        item_data["description_en_is_fallback"] = True
    else:
        desc_pt, desc_en = extract_descriptions_from_docx(docx_file_path)
        item_data["description_pt"] = desc_pt
        item_data["description_en"] = desc_en
        
        if not desc_pt and not desc_en and "Error reading DOCX" not in desc_pt:
            extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Nenhuma descrição (PT ou EN) extraída do description.docx. Usando fallback do CSV se disponível.")
            item_data["description_pt"] = item.get("description_pt_raw_csv", "Descrição em Português não disponível.")
            item_data["description_en"] = item_data["description_pt"]
            item_data["description_en_is_fallback"] = True
        elif not desc_pt and "Error reading DOCX" not in desc_pt:
            extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Descrição em Português ausente no description.docx. Usando fallback do CSV se disponível, ou EN se PT não existir.")
            item_data["description_pt"] = item.get("description_pt_raw_csv", desc_en if desc_en else "Descrição em Português não disponível.")
            if not desc_en: 
                item_data["description_en"] = item_data["description_pt"]
                item_data["description_en_is_fallback"] = True 
        elif not desc_en and "Error reading DOCX" not in desc_en:
            extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Descrição em Inglês ausente no description.docx. Usando descrição em PT como fallback.")
            item_data["description_en"] = item_data["description_pt"] # Fallback EN to PT
            item_data["description_en_is_fallback"] = True
        elif "Error reading DOCX" in desc_pt:
             extraction_issues.append(f"Item '{item_title_for_report}' (Pasta: {folder_name}, Linha CSV: {original_csv_row}): Erro ao ler description.docx: {desc_pt}. Usando fallback do CSV.")
             item_data["description_pt"] = item.get("description_pt_raw_csv", "Erro ao ler descrição.")
             item_data["description_en"] = item_data["description_pt"]
             item_data["description_en_is_fallback"] = True

    full_items_data.append(item_data)

# --- Save Full Data and Extraction Report ---
with open(output_full_data_path, 'w', encoding='utf-8') as outfile:
    json.dump(full_items_data, outfile, indent=4, ensure_ascii=False)

with open(output_extraction_report_path, 'w', encoding='utf-8') as r_file:
    r_file.write("Relatório de Extração de Imagens e Descrições (Pós-Reset)\n")
    r_file.write("===========================================================\n\n")
    r_file.write(f"Total de itens processados para extração: {len(mapped_items)}\n")
    r_file.write(f"Total de itens com dados completos (após extração): {len(full_items_data)}\n\n")
    if extraction_issues:
        r_file.write("Problemas e Avisos durante a extração:\n")
        for issue in extraction_issues:
            r_file.write(f"- {issue}\n")
    else:
        r_file.write("Nenhum problema maior encontrado durante a extração.\n")
    r_file.write("\nNotas:\n- 'description_en_is_fallback': true indica que a descrição em Inglês foi copiada da Portuguesa por falta de uma versão em Inglês explícita no DOCX.\n")
    r_file.write("===========================================================\n")
    r_file.write("Fim do Relatório de Extração.\n")

print(f"Extração de imagens e descrições (pós-reset) concluída. Dados completos salvos em {output_full_data_path}. Relatório em {output_extraction_report_path}")

